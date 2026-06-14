"""
STT 변환 + Word 파일 생성 API
- POST /api/v1/stt/convert/{session_id}  : 특정 녹음 파일 STT 변환
- GET  /api/v1/stt/status/{session_id}   : 변환 상태 확인
- GET  /api/v1/stt/transcripts           : 변환된 Word 파일 목록
- GET  /api/v1/stt/download              : Word 파일 다운로드 URL
"""
from fastapi import APIRouter, BackgroundTasks, HTTPException
from loguru import logger
from datetime import datetime
import os, io, uuid, tempfile, asyncio

router = APIRouter(prefix="/stt", tags=["STT"])

# ── 변환 상태 추적 ────────────────────────────────────────
_stt_jobs: dict = {}
# {session_id: {status, progress, result, error}}

# ── MinIO S3 클라이언트 ───────────────────────────────────
def get_s3():
    import boto3, subprocess
    from botocore.config import Config
    try:
        r = subprocess.run(
            ['docker','inspect','ts_minio','--format',
             '{{range .NetworkSettings.Networks}}{{.IPAddress}} {{end}}'],
            capture_output=True, text=True, timeout=3
        )
        ips = [ip for ip in r.stdout.strip().split() if ip.startswith('172.')]
        ip  = ips[0] if ips else os.environ.get("MINIO_IP","172.31.0.6")
    except Exception:
        ip = os.environ.get("MINIO_IP","172.31.0.6")

    return boto3.client(
        "s3",
        endpoint_url=f"http://{ip}:9000",
        aws_access_key_id=os.environ.get("MINIO_ROOT_USER","minioadmin"),
        aws_secret_access_key=os.environ.get("MINIO_ROOT_PASSWORD","vitnap@ssw0rd"),
        config=Config(signature_version="s3v4",s3={"addressing_style":"path"}),
        region_name="us-east-1",
    )

BUCKET_REC  = "recordings"
BUCKET_TRAN = "transcripts"

# ────────────────────────────────────────────────────────────
# Whisper 모델 로드 (싱글톤)
# ────────────────────────────────────────────────────────────
_whisper_model = None

def load_whisper():
    global _whisper_model
    if _whisper_model is None:
        import whisper
        logger.info("Whisper 모델 로딩 중 (small)...")
        _whisper_model = whisper.load_model("small")
        logger.info("✅ Whisper small 모델 로드 완료")
    return _whisper_model

# ────────────────────────────────────────────────────────────
# 백그라운드 STT 변환 작업
# ────────────────────────────────────────────────────────────
async def run_stt_job(job_id: str, rec_path: str, db_session_id: str,
                      channel: str, title: str):
    """
    1. MinIO에서 음성 파일 다운로드
    2. Whisper STT 변환
    3. Word 파일 생성
    4. Word 파일 MinIO 업로드
    5. DB 업데이트
    """
    _stt_jobs[job_id] = {"status":"running","progress":0,"result":None,"error":None}

    try:
        # ── 1. MinIO에서 음성 파일 다운로드 ─────────────────
        _stt_jobs[job_id]["progress"] = 10
        logger.info(f"[STT] MinIO 다운로드: {rec_path}")
        s3 = get_s3()
        bucket, key = rec_path.split("/", 1)

        with tempfile.NamedTemporaryFile(suffix=".webm", delete=False) as tmp:
            tmp_path = tmp.name
            s3.download_fileobj(bucket, key, tmp)

        logger.info(f"[STT] 다운로드 완료: {tmp_path}")
        _stt_jobs[job_id]["progress"] = 30

        # ── 2. Whisper STT 변환 ──────────────────────────────
        logger.info("[STT] Whisper 변환 시작...")
        model = await asyncio.get_event_loop().run_in_executor(
            None, load_whisper
        )
        result = await asyncio.get_event_loop().run_in_executor(
            None,
            lambda: model.transcribe(
                tmp_path,
                language="ko",          # 한국어 우선
                task="transcribe",
                fp16=False,             # CPU 모드
                verbose=False
            )
        )
        os.unlink(tmp_path)  # 임시 파일 삭제

        transcript_text = result.get("text","").strip()
        segments        = result.get("segments", [])
        logger.info(f"[STT] 변환 완료: {len(transcript_text)}자, {len(segments)}세그먼트")
        _stt_jobs[job_id]["progress"] = 70

        # ── 3. Word 파일 생성 ────────────────────────────────
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 제목
        title_para = doc.add_heading(level=0)
        title_run  = title_para.add_run("📋 통역 방송 텍스트")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 메타 정보
        doc.add_paragraph(f"📅 날짜: {datetime.now().strftime('%Y년 %m월 %d일')}")
        doc.add_paragraph(f"🌐 채널: {channel}")
        doc.add_paragraph(f"📡 방송 제목: {title or '통역 방송'}")
        doc.add_paragraph(f"🤖 변환 엔진: OpenAI Whisper (small)")
        doc.add_paragraph("")  # 빈 줄

        # 구분선
        doc.add_heading("■ 전체 텍스트", level=1)

        # 전체 텍스트
        full_para = doc.add_paragraph(transcript_text)
        full_para.paragraph_format.space_after = Pt(12)

        # 타임스탬프별 세그먼트
        if segments:
            doc.add_heading("■ 타임스탬프별 내용", level=1)
            for seg in segments:
                start = int(seg.get("start", 0))
                end   = int(seg.get("end",   0))
                text  = seg.get("text","").strip()
                h_s, m_s, s_s = start//3600, (start%3600)//60, start%60
                h_e, m_e, s_e = end//3600,   (end%3600)//60,   end%60
                time_str = (f"[{h_s:02d}:{m_s:02d}:{s_s:02d}"
                           f" → {h_e:02d}:{m_e:02d}:{s_e:02d}]")
                p = doc.add_paragraph()
                p.add_run(time_str + " ").bold = True
                p.add_run(text)

        # Word 파일을 메모리에 저장
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        docx_bytes = docx_buffer.getvalue()

        _stt_jobs[job_id]["progress"] = 85
        logger.info(f"[STT] Word 파일 생성 완료: {len(docx_bytes)/1024:.1f}KB")

        # ── 4. Word 파일 MinIO 업로드 ────────────────────────
        now       = datetime.now()
        date_str  = now.strftime("%Y%m%d")
        time_str2 = now.strftime("%H%M%S")
        docx_key  = f"{date_str}/{channel}/transcript_{time_str2}_{job_id[:8]}.docx"

        s3.put_object(
            Bucket=BUCKET_TRAN,
            Key=docx_key,
            Body=docx_bytes,
            ContentType=(
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            )
        )
        transcript_path = f"{BUCKET_TRAN}/{docx_key}"
        logger.info(f"[STT] Word MinIO 저장: {transcript_path}")
        _stt_jobs[job_id]["progress"] = 95

        # ── 5. DB 업데이트 ───────────────────────────────────
        if db_session_id:
            try:
                from app.database import AsyncSessionLocal
                from sqlalchemy import text
                async with AsyncSessionLocal() as db:
                    await db.execute(text(
                        "UPDATE broadcast.sessions"
                        " SET transcript_path=:tp, stt_status='completed',"
                        "     updated_at=NOW()"
                        " WHERE id=:id"
                    ), {"tp": transcript_path, "id": db_session_id})
                    await db.commit()
                logger.info(f"[STT] DB 업데이트 완료: {db_session_id}")
            except Exception as e:
                logger.warning(f"[STT] DB 업데이트 실패: {e}")

        _stt_jobs[job_id].update({
            "status":          "completed",
            "progress":        100,
            "result": {
                "transcript_path": transcript_path,
                "text_length":     len(transcript_text),
                "segments":        len(segments),
                "docx_key":        docx_key,
            }
        })
        logger.info(f"[STT] 작업 완료: {job_id}")

    except Exception as e:
        logger.error(f"[STT] 작업 실패: {job_id} - {e}")
        _stt_jobs[job_id].update({
            "status": "failed",
            "error":  str(e)
        })
        # DB 실패 상태 업데이트
        if db_session_id:
            try:
                from app.database import AsyncSessionLocal
                from sqlalchemy import text
                async with AsyncSessionLocal() as db:
                    await db.execute(text(
                        "UPDATE broadcast.sessions"
                        " SET stt_status='failed', updated_at=NOW()"
                        " WHERE id=:id"
                    ), {"id": db_session_id})
                    await db.commit()
            except Exception:
                pass

# ────────────────────────────────────────────────────────────
# API 엔드포인트
# ────────────────────────────────────────────────────────────

@router.post("/convert")
async def stt_convert(data: dict, background_tasks: BackgroundTasks):
    """
    STT 변환 시작
    body: {rec_path, db_session_id, channel, title}
    """
    rec_path      = data.get("rec_path")
    db_session_id = data.get("db_session_id")
    channel       = data.get("channel", "unknown")
    title         = data.get("title",   "통역 방송")

    if not rec_path:
        raise HTTPException(status_code=400, detail="rec_path 필수")

    job_id = str(uuid.uuid4())[:8]
    _stt_jobs[job_id] = {
        "status":   "queued",
        "progress": 0,
        "result":   None,
        "error":    None
    }

    background_tasks.add_task(
        run_stt_job, job_id, rec_path, db_session_id, channel, title
    )

    logger.info(f"[STT] 변환 작업 등록: {job_id} / {rec_path}")
    return {
        "success": True,
        "job_id":  job_id,
        "message": "STT 변환이 백그라운드에서 시작되었습니다"
    }

@router.get("/status/{job_id}")
async def stt_status(job_id: str):
    """STT 변환 진행 상태 확인"""
    job = _stt_jobs.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="작업 없음")
    return {"success": True, "job_id": job_id, **job}

@router.get("/jobs")
async def stt_jobs():
    """전체 STT 작업 목록"""
    return {"success": True, "jobs": _stt_jobs}

@router.get("/transcripts")
async def stt_transcripts(prefix: str = ""):
    """MinIO transcripts 버킷 Word 파일 목록"""
    try:
        s3  = get_s3()
        res = s3.list_objects_v2(Bucket=BUCKET_TRAN, Prefix=prefix)
        files = []
        for obj in res.get("Contents", []):
            files.append({
                "name":          obj["Key"],
                "size_kb":       round(obj["Size"] / 1024, 1),
                "last_modified": obj["LastModified"].isoformat(),
            })
        return {"success": True, "files": files}
    except Exception as e:
        return {"success": False, "files": [], "error": str(e)}

@router.get("/download-url")
async def stt_download_url(path: str):
    """Word 파일 다운로드 URL (1시간 유효)"""
    try:
        from datetime import timedelta
        s3 = get_s3()
        bucket, key = path.split("/", 1)
        url = s3.generate_presigned_url(
            "get_object",
            Params={"Bucket": bucket, "Key": key},
            ExpiresIn=3600
        )
        return {"success": True, "url": url}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/convert-latest")
async def stt_convert_latest(background_tasks: BackgroundTasks):
    """가장 최근 녹음 파일을 자동으로 STT 변환"""
    try:
        s3  = get_s3()
        res = s3.list_objects_v2(Bucket=BUCKET_REC)
        files = sorted(
            res.get("Contents", []),
            key=lambda x: x["LastModified"],
            reverse=True
        )
        if not files:
            return {"success": False, "message": "녹음 파일 없음"}

        latest = files[0]
        rec_path = f"{BUCKET_REC}/{latest['Key']}"

        job_id = str(uuid.uuid4())[:8]
        _stt_jobs[job_id] = {
            "status": "queued", "progress": 0,
            "result": None, "error": None
        }
        background_tasks.add_task(
            run_stt_job, job_id, rec_path, None, "unknown", "최근 방송"
        )
        return {
            "success":  True,
            "job_id":   job_id,
            "rec_path": rec_path,
            "message":  f"최근 파일({latest['Key']}) STT 변환 시작"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
